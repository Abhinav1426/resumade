import json
import os
from datetime import datetime
import fitz
# import pytesseract
# from PIL import Image
import io
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from fastapi import HTTPException


class FileOperations:
    """File Operations for Resume Builder"""
    def __init__(self):
        self.base_dir = os.path.dirname(os.path.abspath(__file__))


    def load_schema_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            schema_file_contents = json.load(f)
        return schema_file_contents

    def save_json_to_file(self, data , file_name='temp'):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = os.path.join(self.base_dir, '..', 'files',f'{file_name}_file_{timestamp}.json')
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"✅ JSON saved to {filename}")

    def save_pdf_to_file(self, bytes, file_name='temp'):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = os.path.join(self.base_dir, '..', 'files', f'{file_name}_file_{timestamp}.pdf')
        with open(filename, 'wb') as f:
            f.write(bytes)

    def extract_text_from_pdf(self, path):
        doc = fitz.open(path)
        full_text = ""
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()

            if text:
                print(f"[Page {page_num}] Text content found — using regular extraction.")
                full_text += f"\n\n--- Page {page_num} (Text) ---\n{text}"
                links = page.get_links()
                for link in links:
                    if 'uri' in link:
                        uri = link['uri']
                        full_text += f"[Link found on Page {page_num}]: {uri}\n"
            else:
                raise HTTPException(
                    status_code=500,
                    detail=f"No text found on page {page_num}. Consider using OCR for image-based PDFs."
                )
                # print(f"[Page {page_num}] No text found — using OCR.")
                # # Render page as image
                # pix = page.get_pixmap(dpi=300)
                # img = Image.open(io.BytesIO(pix.tobytes("png")))
                #
                # # Perform OCR on the image
                # ocr_text = pytesseract.image_to_string(img)
                # full_text += f"\n\n--- Page {page_num} (OCR) ---\n{ocr_text}"

        doc.close()
        return full_text

    def extract_text_from_pdf_bytes(self, file_bytes):
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = ""
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                print(f"[Page {page_num}] Text content found — using regular extraction.")
                full_text += f"\n\n--- Page {page_num} (Text) ---\n{text}"
                links = page.get_links()
                for link in links:
                    if 'uri' in link:
                        uri = link['uri']
                        full_text += f"[Link found on Page {page_num}]: {uri}\n"
            else:
                print(f"[Page {page_num}] No text found — using OCR.")
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img)
                full_text += f"\n\n--- Page {page_num} (OCR) ---\n{ocr_text}"
        doc.close()
        return full_text

    def extract_text_from_doc(self, path):
        doc = Document(path)
        full_text = ""
        links = []
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        for rel in doc.part.rels.values():
            if rel.reltype == RT.HYPERLINK:
                links.append(rel._target)

        # Append links at the end
        if links:
            full_text += "\n\n--- Hyperlinks Found ---\n"
            for idx, uri in enumerate(links, start=1):
                full_text += f"[Link {idx}]: {uri}\n"

        return full_text

    def extract_text_from_doc_bytes(self, file_bytes):
        doc = Document(io.BytesIO(file_bytes))
        full_text = ""
        links = []
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        for rel in doc.part.rels.values():
            if rel.reltype == RT.HYPERLINK:
                links.append(rel._target)
        if links:
            full_text += "\n\n--- Hyperlinks Found ---\n"
            for idx, uri in enumerate(links, start=1):
                full_text += f"[Link {idx}]: {uri}\n"
        return full_text

    def extract_text_from_file(self, path):
        if '.pdf' in path:
            return self.extract_text_from_pdf(path)
        elif '.doc' in path:
            return self.extract_text_from_doc(path)

    def extract_text_from_file_bytes(self, file_bytes, filename):
        if filename.lower().endswith('.pdf'):
            return self.extract_text_from_pdf_bytes(file_bytes)
        elif filename.lower().endswith('.docx') or filename.lower().endswith('.doc'):
            return self.extract_text_from_doc_bytes(file_bytes)
        else:
            raise ValueError("Unsupported file type for byte extraction.")

    def clean_text(self, text):
        return '\n'.join([line.strip() for line in text.splitlines() if line.strip()])

    def resumefileToText(self,pdf_path):
        text = self.extract_text_from_file(pdf_path)
        text = self.clean_text(text)
        return text

